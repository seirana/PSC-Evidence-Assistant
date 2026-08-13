import hashlib
from pathlib import Path
from typing import List

from docx import Document
from pypdf import PdfReader


# File types that our corpus currently supports
SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


def stable_id(text: str, n: int = 10) -> str:
    """
    Create a stable short ID from a string.

    The same input text will always produce the same ID.
    """
    h = hashlib.sha256(text.encode("utf-8")).hexdigest()
    return h[:n]


def read_text_file(path: Path) -> str:
    """
    Read a plain-text or Markdown file.
    """
    return path.read_text(
        encoding="utf-8",
        errors="ignore",
    )


def read_docx_file(path: Path) -> str:
    """
    Extract text from a .docx Word document.

    This reads:
    - normal paragraphs
    - text inside tables
    """
    document = Document(path)

    parts = []

    # Read normal paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # Read tables
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    cells.append(cell_text)

            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def read_pdf_file(path: Path) -> str:
    """
    Extract text from a PDF.

    Important:
    This works for PDFs that contain actual text.

    It will NOT properly read scanned/image-only PDFs.
    Those would require OCR.
    """
    reader = PdfReader(path)

    pages = []

    for page in reader.pages:
        text = page.extract_text()

        if text:
            pages.append(text.strip())

    return "\n\n".join(pages)


def read_document(path: Path) -> str:
    """
    Read a supported document and return its contents as plain text.

    Supported formats:
    - .txt
    - .md
    - .docx
    - .pdf
    """
    extension = path.suffix.lower()

    if extension in {".txt", ".md"}:
        return read_text_file(path)

    if extension == ".docx":
        return read_docx_file(path)

    if extension == ".pdf":
        return read_pdf_file(path)

    raise ValueError(
        f"Unsupported file type: {extension}"
    )


def list_corpus_files(corpus_dir: Path) -> List[Path]:
    """
    Find all supported documents inside the corpus directory
    and all of its subdirectories.
    """
    return sorted(
        [
            path
            for path in corpus_dir.rglob("*")
            if path.is_file()
            and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ]
    )
